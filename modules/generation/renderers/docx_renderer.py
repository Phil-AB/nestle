"""
DOCX renderer implementation.

This renderer handles Microsoft Word document generation.
It self-registers with the RendererRegistry - NO factory changes needed.
"""

from typing import Dict, Any, Optional, List, TYPE_CHECKING
import time
import uuid
import re
from pathlib import Path

from modules.generation.core.interfaces import IRenderer, GenerationResult
from modules.generation.core.exceptions import RendererException
from modules.generation.core.registry import register_renderer
from modules.generation.config import get_generation_config
from shared.utils.logger import setup_logger

logger = setup_logger(__name__)

# Try to import python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None  # Type hint placeholder
    logger.warning("python-docx not installed. DOCX renderer will not work. Install with: pip install python-docx")


@register_renderer("docx")  # ← SELF-REGISTERS! Zero factory changes.
class DocxRenderer(IRenderer):
    """
    Microsoft Word (DOCX) renderer.
    
    Uses python-docx library to populate Word templates with data.
    """
    
    def __init__(self, config: Dict[str, Any], project_root: Optional[Path] = None):
        """
        Initialize DOCX renderer with configuration.
        
        Args:
            config: Renderer configuration
            project_root: Project root path for resolving relative template paths (optional)
        """
        super().__init__(config)
        
        if not DOCX_AVAILABLE:
            raise RendererException("python-docx library not installed. Install with: pip install python-docx")
        
        # Set project root for path resolution
        gen_config = get_generation_config()
        self.project_root = Path(project_root) if project_root else gen_config.project_root
        self.output_dir = gen_config.output_dir
        
        # Load renderer options from config
        options = config.get('options', {})
        self.preserve_formatting = options.get('preserve_formatting', True)
        self.auto_adjust_tables = options.get('auto_adjust_tables', True)
        self.embed_images = options.get('embed_images', True)
        
        logger.info(f"✅ Initialized DocxRenderer (root: {self.project_root})")
    
    async def render(
        self,
        template_path: str,
        data: Dict[str, Any],
        options: Optional[Dict[str, Any]] = None
    ) -> GenerationResult:
        """
        Render DOCX template with data.
        
        Args:
            template_path: Path to .docx template file
            data: Mapped data to populate template
            options: Optional rendering options
        
        Returns:
            GenerationResult with generated document
        """
        start_time = time.time()
        options = options or {}
        job_id = options.get('job_id', str(uuid.uuid4()))
        
        try:
            logger.info(f"Rendering DOCX template: {template_path}")
            
            # Convert to absolute path if relative
            template_path_obj = Path(template_path)
            if not template_path_obj.is_absolute():
                # Relative to configured project root
                template_path_obj = self.project_root / template_path
            
            # Validate template exists
            if not template_path_obj.exists():
                raise RendererException(f"Template not found: {template_path_obj}")
            
            # Load template
            doc = Document(str(template_path_obj))
            
            # Handle both formats: flat {field: value} or nested {fields: {field: value}}
            if 'fields' in data:
                # Nested format from data provider
                fields = data['fields']
                tables_data = data.get('tables', {})
            else:
                # Flat format from mapper (after mapping)
                # Extract tables and use rest as fields
                tables_data = data.pop('tables', {}) if isinstance(data, dict) else {}
                fields = data
            
            logger.info(f"📄 Replacing {len(fields)} fields in template")
            logger.info(f"📋 Field names: {list(fields.keys())[:10]}")
            logger.info(f"📋 Sample values: {dict(list(fields.items())[:3]) if fields else {}}")
            
            # Replace field placeholders
            if fields:
                logger.info(f"🔄 Calling _replace_fields with {len(fields)} fields")
                self._replace_fields(doc, fields)
                logger.info(f"✅ _replace_fields completed")
            else:
                logger.warning("⚠️  No fields to replace!")
            
            # Populate tables (e.g., line items)
            if tables_data:
                self._populate_tables(doc, tables_data)
            
            # Add images if requested
            if self.embed_images and 'images' in data:
                self._add_images(doc, data['images'])
            
            # Generate output
            output_path = options.get('output_path')
            if not output_path:
                # Use configured output directory
                self.output_dir.mkdir(parents=True, exist_ok=True)
                output_path = str(self.output_dir / f"{job_id}.docx")
            
            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            
            # Save document
            doc.save(output_path)
            
            generation_time = (time.time() - start_time) * 1000  # Convert to ms
            
            logger.info(f"✅ Successfully rendered DOCX in {generation_time:.2f}ms")
            
            return GenerationResult(
                success=True,
                job_id=job_id,
                output_path=output_path,
                output_format="docx",
                template_name=Path(template_path).stem,
                renderer_name=self.renderer_name,
                generation_time_ms=generation_time,
                metadata={
                    "template_path": template_path,
                    "preserve_formatting": self.preserve_formatting,
                }
            )
            
        except Exception as e:
            logger.error(f"❌ Failed to render DOCX: {str(e)}")
            return GenerationResult(
                success=False,
                job_id=job_id,
                error_message=f"DOCX rendering failed: {str(e)}",
                error_details={"template_path": template_path}
            )
    
    def _replace_fields(self, doc: Document, fields: Dict[str, Any]) -> None:
        """Replace {{field_name}} placeholders with actual values."""
        import re
        
        replaced_count = 0
        logger.info(f"🔍 Starting field replacement in {len(doc.paragraphs)} paragraphs")
        
        # Replace in paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            before = paragraph.text
            self._replace_in_paragraph(paragraph, fields)
            after = paragraph.text
            if before != after:
                replaced_count += 1
                logger.debug(f"✏️  Para {i}: '{before[:50]}' → '{after[:50]}'")
        
        logger.info(f"📝 Replaced in {replaced_count} paragraphs")
        
        # Replace in tables
        table_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        before = paragraph.text
                        self._replace_in_paragraph(paragraph, fields)
                        if paragraph.text != before:
                            table_count += 1
        
        logger.info(f"📊 Replaced in {table_count} table cells")
    
    def _replace_in_paragraph(self, paragraph, fields: Dict[str, Any]) -> None:
        """Replace placeholders in a single paragraph preserving formatting."""
        import re
        
        # Get full paragraph text
        full_text = paragraph.text
        
        # Check if any placeholders exist
        has_placeholder = False
        for field_name in fields.keys():
            if f"{{{{{field_name}}}}}" in full_text:
                has_placeholder = True
                break
        
        if not has_placeholder:
            return
        
        # Replace all placeholders in full text
        new_text = full_text
        for field_name, field_value in fields.items():
            placeholder = f"{{{{{field_name}}}}}"
            value_str = str(field_value) if field_value is not None else ""
            new_text = new_text.replace(placeholder, value_str)
        
        # If text changed, update the paragraph
        if new_text != full_text:
            # Clear existing runs
            for run in paragraph.runs:
                run.text = ""
            
            # Add new text in first run (preserves some formatting)
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)
    
    def _populate_tables(self, doc: Document, tables_data: Dict[str, List[Dict]]) -> None:
        """
        Populate tables with array data (e.g., line items).
        
        Looks for placeholders like {{table:line_items}} in the first row,
        then removes that row and adds data rows.
        """
        for table in doc.tables:
            # Check first row for table placeholder
            if not table.rows:
                continue
                
            first_row_text = " ".join([cell.text for cell in table.rows[0].cells])
            
            for table_name, rows_data in tables_data.items():
                placeholder = f"{{{{table:{table_name}}}}}"
                
                if placeholder in first_row_text:
                    logger.debug(f"Found table placeholder for '{table_name}'")
                    
                    # Get column headers from second row (if exists)
                    if len(table.rows) > 1:
                        header_row = table.rows[1]
                        column_names = [cell.text for cell in header_row.cells]
                    else:
                        # Infer from data
                        column_names = list(rows_data[0].keys()) if rows_data else []
                    
                    # Remove placeholder row
                    table._element.remove(table.rows[0]._element)
                    
                    # Add data rows
                    for row_data in rows_data:
                        row = table.add_row()
                        
                        # Populate cells based on column names
                        for idx, column_name in enumerate(column_names):
                            if idx < len(row.cells):
                                value = row_data.get(column_name, "")
                                row.cells[idx].text = str(value) if value is not None else ""
                    
                    logger.debug(f"Populated table '{table_name}' with {len(rows_data)} rows")
    
    def _add_images(self, doc: Document, images: Dict[str, str]) -> None:
        """
        Add images to document.
        
        Looks for placeholders like {{image:company_logo}} and replaces with image.
        """
        for paragraph in doc.paragraphs:
            for image_name, image_path in images.items():
                placeholder = f"{{{{image:{image_name}}}}}"
                if placeholder in paragraph.text:
                    # Validate image exists
                    if not Path(image_path).exists():
                        logger.warning(f"Image not found: {image_path}")
                        continue
                    
                    # Remove placeholder text
                    paragraph.text = paragraph.text.replace(placeholder, "")
                    
                    # Add image
                    run = paragraph.add_run()
                    try:
                        run.add_picture(image_path, width=Inches(2))
                        logger.debug(f"Added image '{image_name}'")
                    except Exception as e:
                        logger.warning(f"Failed to add image '{image_name}': {str(e)}")
    
    def validate_template(self, template_path: str) -> bool:
        """Validate DOCX template."""
        try:
            if not Path(template_path).exists():
                return False
            
            # Try to open the template
            doc = Document(template_path)
            return True
        except Exception as e:
            logger.error(f"Template validation failed: {str(e)}")
            return False
    
    def get_template_fields(self, template_path: str) -> List[str]:
        """Extract field placeholders from template."""
        if not Path(template_path).exists():
            return []
        
        doc = Document(template_path)
        fields = set()
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            # Find {{field_name}} patterns
            matches = re.findall(r'\{\{([^}]+)\}\}', paragraph.text)
            fields.update(matches)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(r'\{\{([^}]+)\}\}', cell.text)
                    fields.update(matches)
        
        return list(fields)
    
    async def health_check(self) -> bool:
        """Check if renderer is healthy."""
        if not DOCX_AVAILABLE:
            return False
        
        try:
            # Try to create a blank document
            doc = Document()
            return True
        except Exception:
            return False
